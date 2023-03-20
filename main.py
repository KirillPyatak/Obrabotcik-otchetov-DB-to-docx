import pandas as pd
#gdown 1si3VrvpwGjPDajRN4-XolfgXahuWE2pk
from docx import Document
import pandas.io.sql as psql
import matplotlib.pyplot as plt
from io import BytesIO
from docx.shared import Pt
from docx.shared import RGBColor
# чтение файла exel
df = pd.read_excel('Таблица ЛС.xlsx')

document = Document() # создается объект

# задаем стиль текста по умолчанию
style = document.styles['Normal']
# название шрифта
style.font.name = 'Times new Roman'
# размер шрифта
style.font.size = Pt(14)
# добавляем первый заголовок
document.add_heading('Результаты выполнения показателей эффективности деятельности профессорского-преподавательского состава Факультета РГГМУ', 1)
p = document.add_paragraph('Отчет о работе: ')

rows, columns = df.shape  # размеры dataframe
table = document.add_table(rows=1, cols=columns) # создаем таблицу
table.style = "Colorful List Accent 1" # определяем стиль

# формируем заголовки таблицы
hdr_cells = table.rows[0].cells
for i in range(columns):
    hdr_cells[i].text = list(df.columns.values)[i]

# заполняем данными из dataframe
for row in range(rows):
    row_cells = table.add_row().cells
    row_data = df.iloc[row].tolist()
    for column in range(columns):
        row_cells[column].text = str(row_data[column])

text = document.add_paragraph('Директор института/Директор факультета')
text.alignment = 0
text.size = Pt(12)

#создаем параграф Дата подпись
paragraph = document.add_paragraph("Дата: ________  Подпись/Расшивровка: _____________/____________ ")
paragraph.alignment = 0
paragraph.size = Pt(12)

paragraph = document.add_paragraph("___________/___________")
paragraph.alignment = 2
paragraph.size = Pt(12)

document.add_page_break()

document.save('test2.docx')

